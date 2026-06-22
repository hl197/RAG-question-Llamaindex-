#!/usr/bin/env python
"""生成项目介绍 Word 文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import datetime


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def main():
    doc = Document()

    # ── 样式设置 ──
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.5
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # ============================================================
    #  封面
    # ============================================================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RAG 智能问答 Agent")
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("基于 LlamaIndex + DeepSeek + ChromaDB 的检索增强生成系统")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(f"项目文档  |  {datetime.date.today().strftime('%Y 年 %m 月')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # ============================================================
    #  目录占位
    # ============================================================
    add_heading(doc, "目录", level=1)
    toc_items = [
        "1. 项目概述",
        "2. 系统架构总览",
        "3. 核心组件详解",
        "    3.1 DeepSeekLLM — 自定义大语言模型封装",
        "    3.2 LocalEmbedding — 纯本地向量化引擎",
        "    3.3 RAGAgent — 智能体协调中枢",
        "    3.4 PersistentChatMemory — 持久化对话记忆",
        "    3.5 Knowledge Loader — 多格式文档解析器",
        "    3.6 ChromaDB Indexer — 向量索引与检索",
        "    3.7 Gradio Frontend — 图形用户界面",
        "4. 数据流与运行流程",
        "    4.1 文件上传 → 知识库入库",
        "    4.2 用户提问 → 智能回答",
        "    4.3 记忆压缩机制",
        "5. 核心技术方法与 Agent 知识点",
        "    5.1 ReAct Agent 模式",
        "    5.2 RAG（检索增强生成）",
        "    5.3 随机投影 Embedding",
        "    5.4 持久化记忆与摘要压缩",
        "6. 项目亮点与面试要点",
        "7. 技术栈与依赖",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.2

    doc.add_page_break()

    # ============================================================
    #  1. 项目概述
    # ============================================================
    add_heading(doc, "1. 项目概述", level=1)
    doc.add_paragraph(
        "RAG 智能问答 Agent 是一个基于检索增强生成（RAG）架构的智能问答系统。"
        "用户可以通过 Web 界面上传 PDF、DOCX、PPTX、TXT、MD、CSV 等多种格式的文档，"
        "系统自动解析、分块、向量化后存入 ChromaDB 向量数据库。"
        "用户提问时，LlamaIndex Agent 自动判断是否需要检索知识库，"
        "从向量库中召回最相关的文档片段，拼接上下文后由 DeepSeek 大语言模型生成回答。"
    )
    doc.add_paragraph(
        "项目完全使用 Python 实现，前端采用 Gradio 框架，"
        "所有 Embedding 向量化过程在本地完成（纯 NumPy，零外部模型下载），"
        "对话历史通过 SQLite 持久化存储，支持跨会话上下文记忆。"
    )

    # ============================================================
    #  2. 系统架构总览
    # ============================================================
    add_heading(doc, "2. 系统架构总览", level=1)
    doc.add_paragraph(
        "系统采用分层架构，自顶向下分为四层："
    )

    arch_text = """
┌─────────────────────────────────────────────────────┐
│                Gradio Web UI (app.py)                │
│   文件上传区  │  会话管理  │  聊天对话窗口           │
└──────────────────────┬──────────────────────────────┘
                       │
┌──────────────────────▼──────────────────────────────┐
│            RAG Agent (rag_agent.py)                  │
│  工具注册  │  记忆管理  │  对话分发                  │
│  ┌──────────────────────────────────────────────┐   │
│  │  ① knowledge_base (RAG检索)                   │   │
│  │  ② knowledge_summary (知识库摘要)             │   │
│  └──────────────────────────────────────────────┘   │
└──────┬──────────────────────────────┬───────────────┘
       │                              │
┌──────▼───────┐           ┌──────────▼────────────┐
│  LLM Layer   │           │   Knowledge Layer     │
│ DeepSeek API │           │  ┌──────────────────┐ │
│ (自定义封装) │           │  │    Loader        │ │
│              │           │  │ PDF/DOCX/PPTX/…  │ │
│              │           │  └────────┬─────────┘ │
│              │           │           │            │
│              │           │  ┌────────▼─────────┐ │
│              │           │  │    Indexer       │ │
│              │           │  │  ChromaDB +      │ │
│              │           │  │  LocalEmbedding  │ │
│              │           │  └──────────────────┘ │
└──────────────┘           └───────────────────────┘
       │                              │
       └──────────┬───────────────────┘
                  │
       ┌──────────▼──────────┐
       │   Memory Layer      │
       │   SQLite 持久化     │
       │   会话管理+摘要压缩 │
       └─────────────────────┘
"""
    add_code(doc, arch_text.strip())

    # ============================================================
    #  3. 核心组件详解
    # ============================================================
    add_heading(doc, "3. 核心组件详解", level=1)

    # --- 3.1 DeepSeekLLM ---
    add_heading(doc, "3.1 DeepSeekLLM — 自定义大语言模型封装", level=2)
    doc.add_paragraph(
        "文件: agent/deepseek_llm.py"
    )
    doc.add_paragraph(
        "DeepSeekLLM 是一个继承自 LlamaIndex LLM 基类的自定义大语言模型封装。"
        "底层使用 OpenAI SDK 调用 DeepSeek API（兼容 OpenAI 协议），"
        "绕开了 LlamaIndex 内置 OpenAI 类的模型名白名单校验。"
    )
    doc.add_paragraph("核心职责：")
    add_bullet(doc, "实现 chat() / complete() 同步对话与补全接口")
    add_bullet(doc, "实现 stream_chat() / stream_complete() 流式接口")
    add_bullet(doc, "实现 achat() / acomplete() 等异步接口")
    add_bullet(doc, "通过 metadata 属性暴露模型信息（模型名、上下文窗口等）")
    add_bullet(doc, "支持 deepseek-chat、deepseek-v4-flash 等任意模型名")
    add_bullet(doc, "Agent 相关知识点：LLM 是 Agent 的「大脑」，负责推理、决策和生成")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("涉及 Agent 知识点：").bold = True
    doc.add_paragraph(
        "LLM (Large Language Model) 是 Agent 系统的推理核心。"
        "在 ReAct 模式中，LLM 负责：理解用户意图 → 决定是否调用工具 → "
        "解析工具返回结果 → 生成最终回答。自定义 LLM 封装展示了"
        "对 LlamaIndex LLM 接口的深入理解，包括同步/异步、流式/非流式的完整实现。"
    )

    # --- 3.2 LocalEmbedding ---
    add_heading(doc, "3.2 LocalEmbedding — 纯本地向量化引擎", level=2)
    doc.add_paragraph(
        "文件: knowledge/local_embedding.py"
    )
    doc.add_paragraph(
        "LocalEmbedding 是一个完全离线的文本向量化引擎，零外部依赖（仅 NumPy）。"
        "它继承了 LlamaIndex 的 BaseEmbedding 基类，实现了标准 Embedding 接口。"
    )
    doc.add_paragraph("核心原理（Random Projection + n-gram 特征）：")
    add_bullet(doc, "字符 n-gram 特征提取：", "特征工程 — ")
    doc.add_paragraph(
        "    对输入文本提取字符级 unigram（单字，权重1）、bigram（二字组，权重2）、"
        "trigram（三字组，权重1）作为特征。这种多粒度特征能同时捕捉词汇和语义信息。"
    )
    add_bullet(doc, "哈希索引：", "哈希技巧 — ")
    doc.add_paragraph(
        "    每个特征通过 MD5 哈希映射到 2^20（约100万）的稀疏空间中，"
        "保证不同特征以极大概率映射到不同位置，避免显式维护字典。"
    )
    add_bullet(doc, "随机投影：", "降维方法 — ")
    doc.add_paragraph(
        "    使用 Achlioptas 稀疏随机投影矩阵（每元素以 2/3 概率为0，"
        "1/6 概率为 +√3，1/6 概率为 -√3），将高维稀疏特征压缩到 256 维稠密向量。"
        "Johnson-Lindenstrauss 引理保证降维后距离关系以高概率保持。"
    )
    add_bullet(doc, "L2 归一化：", "归一化 — ")
    doc.add_paragraph(
        "    最终向量经过 L2 归一化，使余弦相似度等价于内积。"
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("涉及 Agent 知识点：").bold = True
    doc.add_paragraph(
        "Embedding 是 RAG 系统的基石。查询和文档被映射到同一向量空间后，"
        "通过向量相似度检索找到最相关的内容。实现自定义 Embedding 需要理解"
        "BaseEmbedding 的抽象接口（_get_text_embedding、_get_query_embedding、"
        "_get_text_embeddings）以及 Pydantic v2 的模型约束。"
    )

    # --- 3.3 RAGAgent ---
    add_heading(doc, "3.3 RAGAgent — 智能体协调中枢", level=2)
    doc.add_paragraph("文件: agent/rag_agent.py")
    doc.add_paragraph(
        "RAGAgent 是整个系统的核心编排器，整合了 LLM、Embedding、向量数据库、"
        "持久化记忆和工具调用。它使用 LlamaIndex 的 AgentRunner + ReActAgentWorker 构建。"
    )
    doc.add_paragraph("核心组件：")
    add_bullet(doc, "LLM 初始化 — 创建 DeepSeekLLM 实例并注入全局 Settings")
    add_bullet(doc, "Embedding 初始化 — 创建 LocalEmbedding 并注入全局 Settings")
    add_bullet(doc, "工具注册 — 向 Agent 注册 knowledge_base 和 knowledge_summary 两个工具")
    add_bullet(doc, "Agent 创建 — 使用 ReActAgentWorker + AgentRunner 构建可推理的 Agent")
    add_bullet(doc, "记忆管理 — 通过 PersistentChatMemory 管理跨会话历史")
    doc.add_paragraph()
    doc.add_paragraph("注册的两个工具：")
    add_bullet(doc,
        "knowledge_base（QueryEngineTool）：RAG 检索工具，从 ChromaDB 中检索与问题最相关的"
        "文档片段，返回原始文本供 LLM 参考。",
        "Tool 1 — "
    )
    add_bullet(doc,
        "knowledge_summary（FunctionTool）：知识库概览工具，返回已上传文件列表、"
        "文档数量、向量数量等统计信息。",
        "Tool 2 — "
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("涉及 Agent 知识点：").bold = True
    doc.add_paragraph(
        "Agent 是具备工具使用能力的 LLM 系统。在 LlamaIndex 框架中，AgentRunner"
        "负责管理 Agent 的生命周期和记忆，ReActAgentWorker 实现 ReAct 推理循环。"
        "Agent 的核心能力是：观察（Observe）→ 思考（Think）→ 行动（Act）→ 观察结果（Observe）的循环。"
        "通过 FunctionTool 和 QueryEngineTool，Agent 获得了调用外部能力和检索知识的能力。"
    )

    # --- 3.4 PersistentChatMemory ---
    add_heading(doc, "3.4 PersistentChatMemory — 持久化对话记忆", level=2)
    doc.add_paragraph("文件: agent/memory.py")
    doc.add_paragraph(
        "PersistentChatMemory 是一个基于 SQLite 的持久化对话记忆模块，"
        "支持跨会话的对话历史管理、自动摘要压缩和会话标题生成。"
    )
    doc.add_paragraph("数据库设计：")
    add_code(doc,
        "conversations 表              summaries 表\n"
        "├── id (PK)                  ├── id (PK)\n"
        "├── session_id               ├── session_id\n"
        "├── role (user/assistant)    ├── summary_text\n"
        "├── content                  ├── message_count\n"
        "├── session_name             ├── token_count\n"
        "├── metadata                 └── created_at\n"
        "└── created_at"
    )
    doc.add_paragraph("核心功能：")
    add_bullet(doc, "create_session() — 创建新会话（UUID8位），自动设置会话标题")
    add_bullet(doc, "save_message() — 保存单条消息，首次用户消息自动截取前30字做标题")
    add_bullet(doc, "load_history() — 加载会话历史为 ChatMessage 列表，供 Agent 恢复上下文")
    add_bullet(doc, "get_session_list() — 列出所有会话，按最近活跃时间排序")
    add_bullet(doc, "estimate_tokens() — 估算会话 token 数（中文≈2token/字，英文≈0.5token/字）")
    add_bullet(doc, "needs_compression() — 判断是否达到摘要压缩阈值（4000 tokens）")
    add_bullet(doc, "save_summary() — 保存压缩后的摘要，保留历史压缩记录")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("涉及 Agent 知识点：").bold = True
    doc.add_paragraph(
        "记忆（Memory）是 Agent 维持对话连贯性的关键组件。ChatMemoryBuffer 是 Agent 的短期工作记忆，"
        "SQLite 是长期持久化存储。当短期记忆超过 token 限制时，系统自动进行摘要压缩——"
        "将早期对话用 LLM 总结成摘要，保留最近4条完整消息。这是 Agent 记忆管理中的"
        "典型策略，平衡上下文长度与信息完整性。"
    )

    # --- 3.5 Knowledge Loader ---
    add_heading(doc, "3.5 Knowledge Loader — 多格式文档解析器", level=2)
    doc.add_paragraph("文件: knowledge/loader.py")
    doc.add_paragraph(
        "Knowledge Loader 负责将各种格式的文档解析为 LlamaIndex 标准的 Document 对象。"
        "通过文件扩展名自动路由到对应解析器。"
    )
    doc.add_paragraph("支持的格式与解析器：")
    formats = [
        ("PDF", "PyMuPDF4LLM (LlamaMarkdownReader)", "高保真 Markdown 输出，保留表格、标题层级"),
        ("DOCX", "LlamaIndex DocxReader", "纯文本 + 元数据"),
        ("PPTX", "LlamaIndex PptxReader", "每页幻灯片为一个 Document"),
        ("TXT", "原生 Python open()", "自动检测编码（UTF-8/GBK）"),
        ("MD", "原生读取", "Markdown 文本"),
        ("CSV", "LlamaIndex CSVReader", "每行为一个 Document"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "格式"
    hdr[1].text = "解析器"
    hdr[2].text = "特点"
    for fmt, parser, note in formats:
        row = table.add_row().cells
        row[0].text = fmt
        row[1].text = parser
        row[2].text = note

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("涉及 Agent 知识点：").bold = True
    doc.add_paragraph(
        "知识库是 RAG Agent 的信息来源。Loader 的设计体现了「适配器模式」——"
        "不同格式的解析器统一输出为标准的 List[Document]，使得上层 Indexer 无需关心"
        "文件格式差异。这是 Agent 工具链中数据预处理的标准实践。"
    )

    # --- 3.6 ChromaDB Indexer ---
    add_heading(doc, "3.6 ChromaDB Indexer — 向量索引与检索", level=2)
    doc.add_paragraph("文件: knowledge/indexer.py")
    doc.add_paragraph(
        "Indexer 负责管理 ChromaDB 向量数据库的初始化、文档索引构建和相似度检索。"
    )
    doc.add_paragraph("核心功能：")
    add_bullet(doc, "init_chroma() — 初始化 ChromaDB PersistentClient，创建/加载 knowledge_base 集合")
    add_bullet(doc, "build_or_update_index() — 增量索引：分块 → 向量化 → 存入 ChromaDB，自动跳过重复文件")
    add_bullet(doc, "get_query_engine() — 返回 RetrieverQueryEngine，支持 top-k 和相似度阈值配置")
    add_bullet(doc, "list_uploaded_files() — 通过 file_registry.json 管理已上传文件列表")
    add_bullet(doc, "clear_knowledge_base() — 清空知识库，删除所有向量数据和注册表")
    doc.add_paragraph()
    doc.add_paragraph("索引构建流程：")
    add_code(doc,
        "Document → SentenceSplitter(chunk=1024, overlap=200)\n"
        "    → Nodes → LocalEmbedding → 256-dim 向量\n"
        "    → ChromaDB (cosine similarity) → file_registry.json"
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("涉及 Agent 知识点：").bold = True
    doc.add_paragraph(
        "向量存储和检索是 RAG Agent 的「外部知识库」。ChromaDB 作为轻量级向量数据库，"
        "支持余弦相似度检索和持久化存储。检索时，Agent 将用户问题向量化后，"
        "在向量空间中查找最相似的文档片段（top_k=5），这些片段作为上下文注入 LLM 的 prompt。"
        "这是 RAG 的核心机制——用检索增强生成，解决 LLM 知识滞后和幻觉问题。"
    )

    # --- 3.7 Gradio Frontend ---
    add_heading(doc, "3.7 Gradio Frontend — 图形用户界面", level=2)
    doc.add_paragraph("文件: app.py")
    doc.add_paragraph(
        "前端基于 Gradio 6.x 构建，采用双栏布局。左侧为知识库管理和会话管理区，"
        "右侧为对话窗口。所有操作通过回调函数与 RAGAgent 交互。"
    )
    doc.add_paragraph("界面布局：")
    add_bullet(doc, "左侧栏：文件上传（UploadButton）、已上传文件列表、知识库统计、会话切换")
    add_bullet(doc, "右侧栏：聊天窗口（Chatbot，messages 格式）、输入框、发送/清空按钮")
    add_bullet(doc, "页面加载时自动刷新：文件列表、知识库统计、会话列表")
    doc.add_paragraph()
    doc.add_paragraph("主要回调函数：")
    add_bullet(doc, "chat_fn() / respond() — 处理用户消息，调用 Agent 生成回答")
    add_bullet(doc, "upload_file() — 上传并索引文件，返回处理结果")
    add_bullet(doc, "refresh_file_list() / refresh_stats() — 刷新知识库状态")
    add_bullet(doc, "new_session() / switch_session() — 创建/切换会话")

    doc.add_page_break()

    # ============================================================
    #  4. 数据流与运行流程
    # ============================================================
    add_heading(doc, "4. 数据流与运行流程", level=1)

    # --- 4.1 文件上传 ---
    add_heading(doc, "4.1 文件上传 → 知识库入库", level=2)
    add_code(doc,
        "用户选择文件 → Gradio UploadButton\n"
        "  │\n"
        "  ▼\n"
        "app.py: upload_file(file_obj)\n"
        "  → 获取临时文件路径 (file_obj.path)\n"
        "  → RAGAgent.upload_file(file_path)\n"
        "      │\n"
        "      ▼\n"
        "    knowledge/loader.py: load_document(file_path)\n"
        "      → 识别扩展名 → 路由到对应解析器\n"
        "      → PDF: PyMuPDF4LLM (Markdown 高保真)\n"
        "      → DOCX/PPTX/CSV: LlamaIndex 内置 Reader\n"
        "      → TXT/MD: 原生读取\n"
        "      → 输出 List[Document] (含文件名元数据)\n"
        "      │\n"
        "      ▼\n"
        "    knowledge/indexer.py: build_or_update_index(docs, filename)\n"
        "      → 检查 file_registry.json (去重)\n"
        "      → SentenceSplitter 分块 (1024 tokens, overlap 200)\n"
        "      → LocalEmbedding 向量化 (256维)\n"
        "      → ChromaDB 存储 (cosine 距离)\n"
        "      → 更新 file_registry.json\n"
        "      │\n"
        "      ▼\n"
        "    返回: {success, filename, nodes_added}\n"
        "  │\n"
        "  ▼\n"
        "前端显示: ✅ 已上传「张三简历.txt」，新增 X 个文档片段"
    )

    # --- 4.2 用户提问 ---
    add_heading(doc, "4.2 用户提问 → 智能回答", level=2)
    add_code(doc,
        "用户在输入框输入消息 → Enter / 点击发送\n"
        "  │\n"
        "  ▼\n"
        "app.py: respond() → chat_fn() → RAGAgent.chat(message, session_id)\n"
        "  │\n"
        "  ▼ 第一步: 记忆层\n"
        "  ├── session_id 不存在 → 创建新会话 (UUID8)\n"
        "  ├── 首次对话 → 注入 system prompt\n"
        "  │   (角色说明 + 工具使用指引)\n"
        "  ├── 非首次 → 检查 token 是否 > 4000\n"
        "  │   是 → 执行记忆压缩 (见 4.3)\n"
        "  │   否 → 从 SQLite 加载历史到 agent buffer\n"
        "  └── 保存用户消息到 SQLite\n"
        "  │\n"
        "  ▼ 第二步: Agent 推理\n"
        "  self._agent.chat(message)\n"
        "    → ReActAgentWorker 开始 ReAct 循环\n"
        "    │\n"
        "    ├── 思考: 需要检索知识库吗?\n"
        "    ├── 行动: 调用 knowledge_base 工具\n"
        "    │   → 用户问题 → LocalEmbedding → 256维向量\n"
        "    │   → ChromaDB 余弦检索 → top 5 相关片段\n"
        "    │   → 片段文本 + 原始问题 → 拼接 prompt\n"
        "    ├── 观察: 获取检索结果\n"
        "    └── 回答: DeepSeek API 生成最终回答\n"
        "    │\n"
        "    └── (或调用 knowledge_summary → 返回统计信息)\n"
        "  │\n"
        "  ▼ 第三步: 保存 + 返回\n"
        "  ├── 保存 Assistant 回复到 SQLite\n"
        "  └── 返回回复文本 → Gradio 显示\n"
        "  │\n"
        "  ▼\n"
        "前端 Chatbot 追加: user消息 + assistant回复"
    )

    # --- 4.3 记忆压缩 ---
    add_heading(doc, "4.3 记忆压缩机制", level=2)
    add_code(doc,
        "条件触发: estimate_tokens(session_id) > 4000\n"
        "  │\n"
        "  ▼\n"
        "1. 取出该会话所有历史消息\n"
        "2. 分割: 前 N-4 条 → 待压缩, 最后 4 条 → 保留\n"
        "3. 用 DeepSeek 生成中文摘要\n"
        "    prompt: 请对以下对话内容进行简洁的中文摘要...\n"
        "4. SQLite 操作:\n"
        "   a. DELETE 该会话所有旧消息\n"
        "   b. INSERT 摘要 (role=system, 前缀 [历史摘要])\n"
        "   c. INSERT 保留的最近 4 条完整消息\n"
        "5. 记录压缩摘要到 summaries 表\n"
        "  │\n"
        "  ▼\n"
        "结果: token 数大幅减少, 最新对话完整保留"
    )

    doc.add_page_break()

    # ============================================================
    #  5. 核心技术方法与 Agent 知识点
    # ============================================================
    add_heading(doc, "5. 核心技术方法与 Agent 知识点", level=1)

    # --- 5.1 ReAct ---
    add_heading(doc, "5.1 ReAct Agent 模式", level=2)
    doc.add_paragraph(
        "ReAct（Reasoning + Acting）是 LLM Agent 的核心推理范式，"
        "由 Shunyu Yao 等人于 2022 年提出。它将推理（Reasoning）和行动（Acting）"
        "交替进行，让 LLM 在思考过程中可以主动调用外部工具获取信息。"
    )
    doc.add_paragraph("在项目中的实现：")
    add_bullet(doc, "使用 LlamaIndex 的 ReActAgentWorker 作为 Agent 的推理引擎")
    add_bullet(doc, "Agent 接收用户问题后，进入「思考→行动→观察」循环")
    add_bullet(doc, "每一步 Agent 可以决定：调用工具、或直接回答")
    add_bullet(doc, "最大推理步数: 15 步，防止无限循环")
    add_bullet(doc, "verbose 模式打印推理过程（Thought/Action/Observation）")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("面试重点：").bold = True
    doc.add_paragraph(
        "ReAct 模式解决了传统 LLM 的「思考与行动分离」问题。"
        "在 RAG 场景中，Agent 先思考是否需要检索知识库，如果需要则调用检索工具，"
        "观察检索结果后再决定下一步——可以继续检索更多信息，也可以基于已有信息回答。"
        "这种模式比简单的「检索→拼接→回答」流水线更灵活，"
        "Agent 可以根据实际情况动态调整策略。"
    )

    # --- 5.2 RAG ---
    add_heading(doc, "5.2 RAG（检索增强生成）", level=2)
    doc.add_paragraph(
        "RAG（Retrieval-Augmented Generation）是 2020 年由 Lewis 等人提出的"
        "将检索与生成结合的 NLP 范式。核心思想：在 LLM 生成回答前，"
        "先从知识库中检索相关信息作为上下文，减少幻觉并支持知识更新。"
    )
    doc.add_paragraph("在项目中的 RAG 流程：")
    add_bullet(doc, "索引阶段：文档 → 分块 → 向量化 → 存储到向量数据库")
    add_bullet(doc, "检索阶段：问题 → 向量化 → 向量数据库相似度检索 → top-k 片段")
    add_bullet(doc, "生成阶段：检索结果 + 原始问题 + 对话历史 → LLM → 最终回答")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("面试重点：").bold = True
    doc.add_paragraph(
        "RAG 解决了 LLM 的三个核心问题：(1) 知识滞后 — 通过检索最新文档获得时效性信息；"
        "(2) 幻觉 — LLM 必须基于检索结果回答，减少编造可能；"
        "(3) 领域知识 — 无需微调即可让 LLM 处理特定领域内容。"
        "本项目使用 LlamaIndex 框架的 VectorStoreIndex 和 RetrieverQueryEngine 实现标准 RAG。"
    )

    # --- 5.3 Random Projection ---
    add_heading(doc, "5.3 随机投影 Embedding", level=2)
    doc.add_paragraph(
        "这是一个基于 Johnson-Lindenstrauss 引理的降维方法。核心思想："
        "高维空间中的点可以投影到低维空间，同时以高概率保持点之间的距离关系。"
    )
    doc.add_paragraph("数学原理：")
    add_code(doc,
        "对于任意 ε ∈ (0, 1/2) 和 N 个点，存在映射 f: R^n → R^m，其中\n"
        "m = O(log(N) / ε²)，使得对所有点对 (u, v) 有:\n"
        "(1-ε)||uv||² ≤ ||f(u)-f(v)||² ≤ (1+ε)||uv||²"
    )
    doc.add_paragraph("项目实现细节：")
    add_bullet(doc, "Achlioptas 稀疏投影：2/3 概率为 0，1/6 概率为 +√3，1/6 概率为 -√3")
    add_bullet(doc, "特征：字符 n-gram（uni/bi/trigram），权重分别为 1/2/1")
    add_bullet(doc, "哈希：MD5 模 2^20 映射到投影矩阵行索引")
    add_bullet(doc, "维度：256 维（可配置），L2 归一化后用于余弦相似度")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("面试重点：").bold = True
    doc.add_paragraph(
        "这个方案在项目中的特殊之处在于：完全避免了下载预训练 Embedding 模型的需求。"
        "在中国网络环境下，HuggingFace 等模型下载非常困难甚至不可用。"
        "随机投影提供了一种「零下载、零依赖、完全离线」的替代方案，"
        "虽然精度不如预训练模型，但在原型验证和资源受限场景下非常有效。"
        "这是一个典型的工程权衡决策——用精度换取可用性和零部署成本。"
    )

    # --- 5.4 持久化记忆 ---
    add_heading(doc, "5.4 持久化记忆与摘要压缩", level=2)
    doc.add_paragraph(
        "项目的记忆系统分为两层：短期工作记忆（ChatMemoryBuffer）和长期持久化存储（SQLite）。"
    )
    doc.add_paragraph("两层记忆架构：")
    add_bullet(doc, "短期记忆：Agent 内部的 ChatMemoryBuffer，保存当前对话上下文，token 上限可配置")
    add_bullet(doc, "长期记忆：SQLite 数据库持久化所有消息，支持跨会话、跨进程恢复")
    doc.add_paragraph("摘要压缩策略：")
    doc.add_paragraph(
        "当单会话 token 数超过 4000 阈值时触发压缩。压缩策略为：保留最近 4 条完整消息，"
        "将之前的所有消息用 LLM 生成摘要。摘要以 system 消息形式存储在对话历史中。"
        "这样既保留了长对话的关键信息，又控制了 LLM 的上下文长度。"
    )

    doc.add_page_break()

    # ============================================================
    #  6. 项目亮点与面试要点
    # ============================================================
    add_heading(doc, "6. 项目亮点与面试要点", level=1)

    doc.add_paragraph("以下是这个项目在简历和面试中最值得突出的技术亮点：")

    highlights = [
        (
            "完整的 Agent 架构实践",
            "从零搭建了基于 LlamaIndex 的 ReAct Agent 系统，包含工具注册、"
            "记忆管理、多轮对话等完整 Agent 能力。展示了从框架选型到工程落地的全流程。"
        ),
        (
            "RAG 全链路实现",
            "从文档解析（PDF/DOCX/PPTX/TXT/MD/CSV）→ 文本分块 → 向量化 → "
            "向量存储 → 相似度检索 → LLM 增强生成，完整实现了 RAG 的每一环节。"
        ),
        (
            "自定义 Embedding 引擎",
            "手写纯 NumPy 本地 Embedding（随机投影 + n-gram 特征），"
            "零外部模型下载。展示了扎实的数学功底（Johnson-Lindenstrauss 引理）"
            "和工程能力（在受限环境下的创新方案）。"
        ),
        (
            "自定义 LLM 封装",
            "深入理解 LlamaIndex 的 LLM 接口体系，手写 DeepSeekLLM 类"
            "实现完整同步/异步、流式/非流式接口。展示了框架扩展能力和 API 设计理解。"
        ),
        (
            "持久化记忆系统",
            "基于 SQLite 的跨会话记忆管理 + 自动摘要压缩。"
            "解决了长对话的上下文长度瓶颈问题，体现了系统设计中的「权衡」思维。"
        ),
        (
            "全栈工程能力",
            "前端（Gradio）+ 后端（Agent）+ 数据层（ChromaDB + SQLite）+ "
            "API 集成（DeepSeek），完整全栈 AI 应用开发能力。"
        ),
    ]

    for title, desc in highlights:
        p = doc.add_paragraph()
        run = p.add_run(f"✦ {title}")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        doc.add_paragraph(desc)

    # ============================================================
    #  7. 技术栈与依赖
    # ============================================================
    add_heading(doc, "7. 技术栈与依赖", level=1)

    tech_table = [
        ("领域", "技术", "用途"),
        ("AI 框架", "LlamaIndex 0.12+", "Agent 框架、RAG 管线、工具系统"),
        ("大语言模型", "DeepSeek (deepseek-chat)", "对话生成、推理决策"),
        ("Embedding", "自定义 LocalEmbedding (NumPy)", "文本向量化，零下载"),
        ("向量数据库", "ChromaDB + cosine", "文档向量存储与相似度检索"),
        ("前端", "Gradio 6.x", "Web 交互界面"),
        ("持久化", "SQLite", "会话历史存储"),
        ("文档解析", "PyMuPDF4LLM", "PDF → Markdown 高保真解析"),
        ("文档解析", "python-docx / python-pptx", "DOCX / PPTX 解析"),
        ("工具", "python-dotenv", "环境变量管理"),
        ("加密", "cryptography (Fernet)", "API Key 加密存储"),
    ]

    table = doc.add_table(rows=len(tech_table), cols=3)
    table.style = "Light Grid Accent 1"
    for i, (area, tech, usage) in enumerate(tech_table):
        row = table.rows[i].cells
        row[0].text = area
        row[1].text = tech
        row[2].text = usage

    doc.add_paragraph()
    add_heading(doc, "项目结构", level=2)
    add_code(doc,
        "RAG-question/\n"
        "├── app.py                      # Gradio 前端入口\n"
        "├── config.py                   # 全局配置 + API Key 解密\n"
        "├── agent/\n"
        "│   ├── deepseek_llm.py         # 自定义 DeepSeek LLM 封装\n"
        "│   ├── rag_agent.py            # Agent 核心：工具/记忆/对话\n"
        "│   └── memory.py              # SQLite 持久化记忆\n"
        "├── knowledge/\n"
        "│   ├── loader.py              # 多格式文档解析器\n"
        "│   ├── indexer.py            # ChromaDB 索引/检索\n"
        "│   └── local_embedding.py    # 纯 NumPy 本地 Embedding\n"
        "├── config/\n"
        "│   └── keys.enc               # 加密后的 API Key\n"
        "├── scripts/\n"
        "│   ├── encrypt_key.py         # API Key 加密工具\n"
        "│   └── generate_report.py    # 本文档生成脚本\n"
        "├── test_data/                 # 测试文件\n"
        "├── data/                      # 运行时数据（gitignored）\n"
        "├── ERROR_LOG.md               # 错误记录\n"
        "└── requirements.txt           # 依赖清单"
    )

    # ── 保存 ──
    output_path = "项目介绍_RAG智能问答Agent.docx"
    doc.save(output_path)
    print(f"✅ 已生成: {output_path}")

if __name__ == "__main__":
    main()
